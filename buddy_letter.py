#!/usr/bin/env python3
# Copyright (c) 2026 Oliver Tillinghast.
# Licensed under the PolyForm Noncommercial License 1.0.0.
# Free for veterans, VSOs, nonprofits and government. Commercial use
# requires a separate licence -- see LICENSE, NOTICE.md, COMMERCIAL.md.
"""
buddy_letter.py — editable buddy-letter (lay statement) templates.

A buddy letter is a statement from someone who observed the member — a
spouse, roommate, supervisor, or fellow airman — describing what they saw.
Its value is in firsthand observation of things a medical record doesn't
capture: what changed, when, and how it showed up day to day.

So these templates deliberately do NOT draft the letter's body. The tool
has no basis to write what someone else witnessed, and a generated
narrative signed by a real person as their own account would be a
fabricated record. What it does instead: pre-fill the administrative
header, state what the medical record already documents (so the writer
knows what they're corroborating and doesn't contradict it by accident),
and give a prompt sheet of what to actually address.

One .docx per confirmed condition, plus a general one.
"""

import os
import re

from docx import Document
from docx.shared import Pt

PROMPTS = [
    "How do you know the veteran, and for how long?",
    "What did you personally see or hear? Describe specific moments, not "
    "general impressions.",
    "When did you first notice it? What was different before versus after?",
    "How does it affect their day-to-day life — work, sleep, relationships, "
    "physical activity?",
    "Has it gotten better, worse, or stayed the same over time?",
    "Anything else you witnessed that a doctor's notes wouldn't capture?",
]

GUIDANCE = [
    "Write in your own words. Plain language is better than formal language.",
    "Only write what you personally witnessed. Don't guess at diagnoses, "
    "causes, or what the VA should decide — firsthand observation is exactly "
    "what makes this statement useful.",
    "Be specific with dates and examples wherever you can, even approximate "
    "ones (\"around summer 2022\" is more useful than \"a while ago\").",
    "If you don't remember something clearly, say so rather than filling it in.",
    "Sign and date it. Include your contact information above.",
]


def _safe_filename(text):
    slug = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
    return slug[:60] or "condition"


def _add_heading(doc, text, size=16):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p


def _add_label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


def build_letter(member_name, condition=None):
    """One buddy-letter template. `condition` None means a general letter."""
    doc = Document()

    subject = condition["condition"] if condition else "general statement"
    _add_heading(doc, "Buddy Letter / Lay Statement")
    doc.add_paragraph(f"Regarding: {member_name}")
    doc.add_paragraph(f"Subject: {subject}")
    doc.add_paragraph()

    _add_label(doc, "Your information (fill in):")
    for label in ("Full name", "Relationship to the veteran",
                  "Address", "Phone", "Email"):
        doc.add_paragraph(f"{label}: ______________________________")
    doc.add_paragraph()

    if condition and condition.get("self_reported"):
        # This is the letter that actually decides something. A documented
        # condition has a clinician's note carrying it; this one has nothing
        # but the member's word, and a buddy who saw it is the only evidence
        # that will ever exist. Say so plainly, so the person writing it
        # understands they are not producing a formality.
        _add_label(doc, "Why this letter matters more than the others:")
        doc.add_paragraph(
            f"There is no medical record for {condition['condition'].lower()}. "
            f"It was never written down at the time — which is ordinary, and "
            f"is not evidence it did not happen. What you personally saw is "
            f"the evidence. Be specific: dates, places, what changed, what "
            f"they stopped being able to do.")
        doc.add_paragraph(
            "Do not guess at a diagnosis or use medical words. Describe what "
            "you observed and let the doctors name it.")
        doc.add_paragraph()
    elif condition:
        _add_label(doc, "What the medical record already documents:")
        icd = f" ({condition['icd10']})" if condition.get("icd10") else ""
        span = (condition["first_seen"] if condition["first_seen"] == condition["last_seen"]
                else f"{condition['first_seen']} to {condition['last_seen']}")
        doc.add_paragraph(f"{condition['condition']}{icd} — documented {span}, "
                          f"{condition['encounters']} encounter"
                          f"{'s' if condition['encounters'] != 1 else ''}.")
        doc.add_paragraph(
            "This is context so your statement lines up with the record. "
            "Don't repeat it back — write what you saw.")
        doc.add_paragraph()

    _add_label(doc, "What to address:")
    for prompt in PROMPTS:
        doc.add_paragraph(prompt, style="List Bullet")
    doc.add_paragraph()

    _add_label(doc, "Your statement:")
    for _ in range(3):
        doc.add_paragraph("")
        doc.add_paragraph("_" * 78)
    doc.add_paragraph()

    doc.add_paragraph("Signature: ____________________________    "
                      "Date: ________________")
    doc.add_paragraph()

    _add_label(doc, "How to write a useful statement:")
    for tip in GUIDANCE:
        doc.add_paragraph(tip, style="List Bullet")

    return doc


def write_letters(member_name, conditions, out_dir):
    """One template per condition plus a general one. Returns paths.

    Self-reported conditions are emitted first and named so they sort first.
    The tool used to generate letters only for conditions already in the
    record -- exactly the ones that need a lay statement least, since they
    already have a clinician's note behind them. The undocumented ones are
    what buddy letters exist for.
    """
    conditions = sorted(conditions,
                        key=lambda c: (not c.get("self_reported"),
                                       c.get("condition", "")))
    os.makedirs(out_dir, exist_ok=True)
    paths = []

    general = os.path.join(out_dir, "buddy_letter_GENERAL.docx")
    build_letter(member_name).save(general)
    paths.append(general)

    for cond in conditions:
        name = _safe_filename(cond["condition"])
        # The undocumented ones are the letters that carry weight, so they
        # sort to the top of the folder instead of alphabetically into the
        # middle of a pile the member may never open.
        prefix = "NEEDED_" if cond.get("self_reported") else ""
        path = os.path.join(out_dir, f"buddy_letter_{prefix}{name}.docx")
        build_letter(member_name, cond).save(path)
        paths.append(path)

    return paths
